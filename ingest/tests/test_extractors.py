"""Tests for ingest.extractors.

PDF and DOCX fixtures are constructed in-test from the same libraries
the production extractors use, which keeps the repo free of binary
fixture blobs. The MD and TXT cases are plain string content. The
"real-world PDF" case proves PDF extraction against the canonical
Document Retention Policy fixture committed under internal/.
"""
from __future__ import annotations

import io
import re
from pathlib import Path

import pytest

from ingest.extractors import (
    DocxExtractor,
    Extractor,
    PdfExtractor,
    TextExtractor,
    UnsupportedFormatError,
    extract,
)

REPO_ROOT = Path(__file__).resolve().parents[2]
RETENTION_PDF = REPO_ROOT / "internal" / "Document Retention Policy.pdf"


@pytest.fixture
def tiny_pdf(tmp_path: Path) -> Path:
    """Build a minimal one-page PDF with the literal text 'Hello PDF World'.

    The source bytes use a hand-rolled PDF skeleton with a bogus xref
    offset; we roundtrip through PdfWriter to emit a clean, warning-
    free PDF that pypdf can parse without complaint.
    """
    import pypdf
    from pypdf import PdfReader, PdfWriter

    src = (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]"
        b"/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
        b"4 0 obj<</Length 44>>stream\n"
        b"BT /F1 24 Tf 100 700 Td (Hello PDF World) Tj ET\n"
        b"endstream endobj\n"
        b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
        b"xref\n0 6\n"
        b"0000000000 65535 f\n"
        b"0000000009 00000 n\n"
        b"0000000052 00000 n\n"
        b"0000000099 00000 n\n"
        b"0000000189 00000 n\n"
        b"0000000276 00000 n\n"
        b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n339\n%%EOF\n"
    )
    reader = PdfReader(io.BytesIO(src))
    writer = PdfWriter(clone_from=reader)
    pdf_path = tmp_path / "tiny.pdf"
    with pdf_path.open("wb") as fh:
        writer.write(fh)
    return pdf_path


@pytest.fixture
def tiny_docx(tmp_path: Path) -> Path:
    """Build a minimal .docx with two paragraphs using python-docx."""
    import docx

    document = docx.Document()
    document.add_paragraph("Hello DOCX World")
    document.add_paragraph("Second paragraph.")
    docx_path = tmp_path / "tiny.docx"
    document.save(str(docx_path))
    return docx_path


@pytest.fixture
def tiny_md(tmp_path: Path) -> Path:
    p = tmp_path / "tiny.md"
    p.write_text("# Heading\n\nA paragraph with **bold** text.", encoding="utf-8")
    return p


@pytest.fixture
def tiny_txt(tmp_path: Path) -> Path:
    p = tmp_path / "tiny.txt"
    p.write_text("plain text content\nline two\n", encoding="utf-8")
    return p


# ---------- per-format dispatch -----------------------------------------------


def test_extract_pdf_returns_text(tiny_pdf: Path):
    text = extract(tiny_pdf)
    assert text.strip() == "Hello PDF World"


def test_extract_docx_returns_text(tiny_docx: Path):
    text = extract(tiny_docx)
    assert "Hello DOCX World" in text
    assert "Second paragraph." in text


def test_extract_md_returns_text(tiny_md: Path):
    text = extract(tiny_md)
    assert text.startswith("# Heading")
    assert "**bold**" in text


def test_extract_txt_returns_text(tiny_txt: Path):
    text = extract(tiny_txt)
    assert "plain text content" in text
    assert "line two" in text


# ---------- case-insensitivity ------------------------------------------------


def test_extract_uppercase_suffix_dispatches(tmp_path: Path):
    """File suffix matching is case-insensitive (.PDF, .DOCX, .MD, .TXT)."""
    p = tmp_path / "shouty.TXT"
    p.write_text("yelling text", encoding="utf-8")
    assert extract(p) == "yelling text"


# ---------- error contracts ---------------------------------------------------


def test_extract_unknown_suffix_raises_unsupported_format(tmp_path: Path):
    p = tmp_path / "foo.xyz"
    p.write_text("does not matter", encoding="utf-8")
    with pytest.raises(UnsupportedFormatError, match=r"\.xyz"):
        extract(p)


def test_extract_unknown_suffix_does_not_require_existence():
    """The format check fires before any filesystem check."""
    with pytest.raises(UnsupportedFormatError, match=r"\.xyz"):
        extract(Path("does_not_exist.xyz"))


def test_extract_missing_supported_file_raises_filenotfound(tmp_path: Path):
    missing = tmp_path / "missing.pdf"
    with pytest.raises(FileNotFoundError, match=re.escape(str(missing))):
        extract(missing)


def test_extract_accepts_str_path(tmp_path: Path):
    """extract() coerces str inputs to Path for caller convenience."""
    p = tmp_path / "as_str.md"
    p.write_text("from a string path", encoding="utf-8")
    assert extract(str(p)).strip() == "from a string path"


# ---------- end-to-end PDF against the real retention policy ------------------


@pytest.mark.skipif(
    not RETENTION_PDF.exists(),
    reason="internal/Document Retention Policy.pdf fixture absent",
)
def test_extract_real_retention_pdf_contains_expected_phrase():
    """AC #2: real-world PDF returns a non-empty string with diocese-specific text."""
    text = extract(RETENTION_PDF)
    assert text.strip()
    assert (
        "Diocese of Pensacola-Tallahassee" in text
        or "Document Retention Policy" in text
    )


# ---------- ABC discoverability -----------------------------------------------


def test_extractors_implement_extractor_abc():
    """Each concrete extractor exposes its extensions and is an Extractor."""
    for cls in (PdfExtractor, DocxExtractor, TextExtractor):
        instance = cls()
        assert isinstance(instance, Extractor)
        assert isinstance(cls.extensions, tuple)
        assert all(ext.startswith(".") and ext == ext.lower() for ext in cls.extensions)
